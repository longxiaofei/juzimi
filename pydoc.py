import pymongo
from docx import Document

client = pymongo.MongoClient('localhost')
db = client['juzimi']
all_collection = db.collection_names()
for collection in all_collection:
	document = Document()
	for data in db[collection].find():
		all_content = data['post_content'] + '            --------- ' + data['post_title']
		document.add_paragraph('').add_run(all_content).bold = True
	document.save(collection+'.doc')
	